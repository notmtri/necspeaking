import base64
import io
import json
import mimetypes
import os
import re
from datetime import datetime

import httpx
from pydub import AudioSegment
from werkzeug.utils import secure_filename


ALLOWED_EXTENSIONS = {'wav', 'mp3', 'm4a', 'webm', 'ogg'}
TRANSCRIPTION_MODEL = os.getenv('GROQ_TRANSCRIPTION_MODEL', 'whisper-large-v3')
GEMINI_GRADING_MODEL = os.getenv('GEMINI_GRADING_MODEL', 'gemini-3.5-flash')
GROQ_GRADING_FALLBACK_MODEL = os.getenv('GROQ_GRADING_FALLBACK_MODEL', 'openai/gpt-oss-120b')
GEMINI_API_URL = 'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent'


GRADING_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "scores": {
            "type": "object",
            "properties": {
                "content": {"type": "number"},
                "accuracy": {"type": "number"},
                "delivery": {"type": "number"},
                "total": {"type": "number"},
            },
            "required": ["content", "accuracy", "delivery", "total"],
        },
        "feedback": {
            "type": "object",
            "properties": {
                "content": {"type": "string"},
                "accuracy": {"type": "string"},
                "delivery": {"type": "string"},
            },
            "required": ["content", "accuracy", "delivery"],
        },
        "sample_response": {"type": "string"},
    },
    "required": ["scores", "feedback", "sample_response"],
}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def cleanup_old_files(upload_folder):
    try:
        current_time = datetime.now().timestamp()
        for filename in os.listdir(upload_folder):
            filepath = os.path.join(upload_folder, filename)
            if filename in ['samples', 'simulations', 'questions.json', 'metadata.json', 'jobs']:
                continue
            if os.path.isfile(filepath):
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > 3600:
                    os.remove(filepath)
    except Exception as error:
        print(f"Cleanup error: {error}")


def get_audio_duration(file_path):
    audio = AudioSegment.from_file(file_path)
    return len(audio) / 1000.0


def convert_to_wav(input_path, output_path):
    audio = AudioSegment.from_file(input_path)
    audio = audio.set_frame_rate(16000).set_channels(1)
    audio.export(output_path, format="wav")
    return output_path


def transcribe_audio(groq_client, file_path):
    with open(file_path, 'rb') as audio_file:
        transcription = groq_client.audio.transcriptions.create(
            file=("audio.wav", audio_file.read()),
            model=TRANSCRIPTION_MODEL,
            response_format="json",
        )

    duration = get_audio_duration(file_path)
    transcript_text = transcription.text if hasattr(transcription, 'text') else str(transcription)

    return {
        "text": transcript_text,
        "words": [],
        "duration": duration
    }


def build_grading_prompt(topic, transcript_data, audio_attached=False):
    transcript_text = transcript_data["text"]
    total_words = len(transcript_text.split())
    duration = transcript_data["duration"]
    words_per_minute = (total_words / duration * 60) if duration > 0 else 0
    audio_instruction = (
        "The original speech audio is attached. Use it to evaluate pronunciation, intonation, "
        "pauses, fluency, confidence, and delivery. Use the transcript for content, vocabulary, "
        "grammar, and examples."
        if audio_attached
        else "Only the transcript and timing metrics are available. Do not claim to hear pronunciation directly."
    )

    return f"""You are an expert English speaking examiner. Grade the following speech response based on this rubric:

**Rubric (Total: 2.0 points)**
1. Content (0.9/2.0 points)
   - Sufficiently address all requirements of the test question
   - Develop supporting ideas with relevant reasons and examples
   - Display a range of original and practical ideas

2. Accuracy (0.6/2.0 points)
   - Demonstrate a wide variety of vocabulary and grammatical structures
   - Make correct use of words, grammatical structures and linking devices
   - Demonstrate correct pronunciation with appropriate intonation

3. Delivery (0.5/2.0 points)
   - Maintain fluency throughout
   - Demonstrate effective use of presentation skills

**Topic/Question:** {topic}

**Speech Transcript:** {transcript_text}

**Speech Metrics:**
- Total words: {total_words}
- Duration: {duration:.1f} seconds
- Speaking pace: {words_per_minute:.0f} words/minute

**Audio Availability:** {audio_instruction}

**Instructions:**
1. Provide scores for each criterion (rounded to 2 decimal places)
2. Give detailed feedback for each criterion with specific examples from the transcript and audio when available
3. Point out both strengths and areas for improvement
4. Generate a comprehensive sample 2.0/2.0 response to the same topic that would take approximately 5 minutes to speak (around 600-750 words). The sample should:
   - Start with "My question is... (if question number is provided), and the prompt is... Here is my response." and then answer the question fully
   - End with "This is the end of my speech. Thank you."
   - Be detailed and well-structured with clear introduction, body paragraphs, and conclusion
   - Include specific examples, explanations, and supporting details
   - Demonstrate sophisticated vocabulary and varied sentence structures
   - Show natural flow with appropriate transitions
   - Be comprehensive enough to fill a 5-minute speaking time
   - Be creative in the introduction to hook the listener's attention
   - Grade at C2 level of the CEFR framework
   - Be very strict

Return only valid JSON matching this shape:
{{
    "scores": {{
        "content": 0.00,
        "accuracy": 0.00,
        "delivery": 0.00,
        "total": 0.00
    }},
    "feedback": {{
        "content": "Detailed feedback with examples...",
        "accuracy": "Detailed feedback with examples...",
        "delivery": "Detailed feedback with examples..."
    }},
    "sample_response": "A complete 2.0/2.0 sample response to the topic..."
}}"""


def parse_grading_json(result_text):
    if "```json" in result_text:
        result_text = result_text.split("```json")[1].split("```")[0].strip()
    elif "```" in result_text:
        result_text = result_text.split("```")[1].split("```")[0].strip()

    result_text = result_text.replace('"', '"').replace('"', '"')
    result_text = result_text.replace("''", "'").replace("''", "'")
    result_text = result_text.replace('\u2018', "'").replace('\u2019', "'")
    result_text = result_text.replace('\u201c', '"').replace('\u201d', '"')
    result_text = result_text.replace('\u2013', '-').replace('\u2014', '-')
    result_text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2060\uFEFF]', '', result_text)
    result_text = result_text.replace('\u202f', ' ')
    result_text = result_text.replace('\ufeff', '')
    result_text = result_text.replace('\u00A0', ' ')
    result_text = re.sub(r'[^\x00-\x7F]+', '', result_text)
    result_text = re.sub(r'[\x00-\x1F\x7F]', '', result_text)

    return json.loads(result_text)


def grade_speech_with_gemini(topic, transcript_data, audio_path):
    gemini_api_key = os.getenv('GEMINI_API_KEY', '').strip()
    if not gemini_api_key:
        raise RuntimeError("GEMINI_API_KEY is not configured.")

    parts = [{"text": build_grading_prompt(topic, transcript_data, audio_attached=bool(audio_path))}]
    if audio_path:
        if audio_path.lower().endswith('.wav'):
            mime_type = 'audio/wav'
        else:
            mime_type = mimetypes.guess_type(audio_path)[0] or 'audio/wav'
        with open(audio_path, 'rb') as audio_file:
            parts.append({
                "inline_data": {
                    "mime_type": mime_type,
                    "data": base64.b64encode(audio_file.read()).decode('ascii'),
                }
            })

    payload = {
        "contents": [{"parts": parts}],
        "generationConfig": {
            "temperature": 0.2,
            "responseFormat": {
                "text": {
                    "mimeType": "application/json",
                    "schema": GRADING_RESPONSE_SCHEMA,
                }
            },
        },
    }

    url = GEMINI_API_URL.format(model=GEMINI_GRADING_MODEL)
    with httpx.Client(timeout=120) as client:
        response = client.post(
            url,
            headers={
                "Content-Type": "application/json",
                "x-goog-api-key": gemini_api_key,
            },
            json=payload,
        )
        response.raise_for_status()
        data = response.json()

    candidates = data.get("candidates") or []
    if not candidates:
        raise RuntimeError("Gemini returned no grading candidate.")

    response_parts = (candidates[0].get("content") or {}).get("parts") or []
    result_text = ''.join(part.get("text", "") for part in response_parts).strip()
    if not result_text:
        raise RuntimeError("Gemini returned an empty grading response.")

    return parse_grading_json(result_text)


def grade_speech_with_groq(groq_client, topic, transcript_data):
    prompt = build_grading_prompt(topic, transcript_data, audio_attached=False)
    response = groq_client.chat.completions.create(
        model=GROQ_GRADING_FALLBACK_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )

    return parse_grading_json(response.choices[0].message.content)


def grade_speech(groq_client, topic, transcript_data, audio_path=''):
    try:
        return grade_speech_with_gemini(topic, transcript_data, audio_path)
    except Exception as error:
        print(f"[ANALYSIS] Gemini grading unavailable, falling back to Groq: {error}")
        return grade_speech_with_groq(groq_client, topic, transcript_data)


def generate_docx(topic, transcript, grading_result):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    title = doc.add_heading('necs. - Speech Feedback Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph()

    doc.add_heading('Score Summary', 1)
    scores = grading_result['scores']

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'

    score_data = [
        ('Content', f"{scores['content']}/0.9"),
        ('Accuracy', f"{scores['accuracy']}/0.6"),
        ('Delivery', f"{scores['delivery']}/0.5"),
        ('', ''),
        ('TOTAL SCORE', f"{scores['total']}/2.0")
    ]

    for index, (criterion, score) in enumerate(score_data):
        table.rows[index].cells[0].text = criterion
        table.rows[index].cells[1].text = score
        if index == 4:
            for cell in table.rows[index].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()

    doc.add_heading('Detailed Feedback', 1)
    feedback = grading_result['feedback']

    doc.add_heading('1. Content', 2)
    doc.add_paragraph(feedback['content'])

    doc.add_heading('2. Accuracy', 2)
    doc.add_paragraph(feedback['accuracy'])

    doc.add_heading('3. Delivery', 2)
    doc.add_paragraph(feedback['delivery'])

    doc.add_page_break()

    doc.add_heading('Your Speech Transcript', 1)
    doc.add_paragraph(transcript)

    doc.add_page_break()

    doc.add_heading('Sample 2.0/2.0 Response', 1)
    doc.add_paragraph(grading_result['sample_response'])

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


def build_job_storage_path(upload_folder, filename):
    jobs_folder = os.path.join(upload_folder, 'jobs')
    os.makedirs(jobs_folder, exist_ok=True)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    secured = secure_filename(filename)
    return os.path.join(jobs_folder, f"{timestamp}_{secured}")
