from flask import Flask, request, jsonify, send_file, send_from_directory, session
from flask_cors import CORS
from dotenv import load_dotenv
import os
import warnings
import base64
import secrets
from functools import wraps
from datetime import datetime, timedelta

warnings.filterwarnings("ignore", message="Core Pydantic V1 functionality")

from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash
from groq import Groq
import json
from pydub import AudioSegment
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import random

from database import db, Question, Sample
import cloudinary
import cloudinary.uploader

load_dotenv()

app = Flask(__name__, static_folder='build', static_url_path='')

# REMOVED REDIS LIMITER - Use custom rate limiting instead
# If you need Redis later, add it back with proper configuration

# Security Configuration
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', secrets.token_hex(32))
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SECURE'] = False  # Set to True if using HTTPS
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=2)

# Database Configuration
database_url = os.getenv('DATABASE_URL')
if database_url and database_url.startswith('postgres://'):
    database_url = database_url.replace('postgres://', 'postgresql://', 1)

app.config['SQLALCHEMY_DATABASE_URI'] = database_url or 'sqlite:///necs.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db.init_app(app)

with app.app_context():
    db.create_all()
    print("✅ Database tables created successfully!")

# Cloudinary Configuration
cloudinary.config(
    cloud_name=os.getenv('CLOUDINARY_CLOUD_NAME'),
    api_key=os.getenv('CLOUDINARY_API_KEY'),
    api_secret=os.getenv('CLOUDINARY_API_SECRET'),
    secure=True
)

# SECURE CORS - Replace * with your actual frontend domain
ALLOWED_ORIGINS = os.getenv('ALLOWED_ORIGINS', 'http://localhost:3000').split(',')

CORS(app, resources={
    r"/api/*": {
        "origins": ALLOWED_ORIGINS,
        "methods": ["GET", "POST", "PUT", "DELETE"],
        "allow_headers": ["Content-Type"],
        "supports_credentials": True
    }
})

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.config['UPLOAD_FOLDER'] = 'uploads'
ALLOWED_EXTENSIONS = {'wav', 'mp3', 'm4a', 'webm', 'ogg'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

groq_client = Groq(api_key=os.getenv("GROQ_API_KEY"))

# ADMIN PASSWORD - FIXED
ADMIN_PASSWORD_HASH = os.getenv('ADMIN_PASSWORD_HASH')
if not ADMIN_PASSWORD_HASH:
    print("⚠️ WARNING: Using fallback password hash. Set ADMIN_PASSWORD_HASH in production!")
    ADMIN_PASSWORD_HASH = generate_password_hash('040108Minhtri')
    print(f"✅ Generated hash for testing: {ADMIN_PASSWORD_HASH[:50]}...")

# Rate Limiting Storage (simple in-memory)
rate_limit_storage = {}

def rate_limit(max_requests=10, window_seconds=60):
    """Simple rate limiting decorator"""
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            client_ip = request.remote_addr
            current_time = datetime.now()
            
            if client_ip not in rate_limit_storage:
                rate_limit_storage[client_ip] = []
            
            # Clean old requests
            rate_limit_storage[client_ip] = [
                req_time for req_time in rate_limit_storage[client_ip]
                if (current_time - req_time).seconds < window_seconds
            ]
            
            if len(rate_limit_storage[client_ip]) >= max_requests:
                return jsonify({"error": "Rate limit exceeded. Try again later."}), 429
            
            rate_limit_storage[client_ip].append(current_time)
            return f(*args, **kwargs)
        return wrapped
    return decorator

def require_admin():
    """Decorator to require admin authentication"""
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            if not session.get('admin_authenticated'):
                return jsonify({"error": "Unauthorized. Admin login required."}), 401
            return f(*args, **kwargs)
        return wrapped
    return decorator

# ============= AUTHENTICATION ROUTES =============

@app.route('/api/admin/login', methods=['POST'])
@rate_limit(max_requests=5, window_seconds=300)
def admin_login():
    """Secure admin login endpoint"""
    try:
        data = request.get_json()
        password = data.get('password', '')
        
        print(f"🔐 Login attempt - Password received: {bool(password)}")
        print(f"🔐 Hash exists: {bool(ADMIN_PASSWORD_HASH)}")
        
        if check_password_hash(ADMIN_PASSWORD_HASH, password):
            session['admin_authenticated'] = True
            session.permanent = True
            print("✅ Login successful!")
            return jsonify({
                "success": True,
                "message": "Login successful"
            })
        else:
            print("❌ Password check failed")
            return jsonify({"error": "Invalid password"}), 401
            
    except Exception as e:
        print(f"❌ Login error: {str(e)}")
        return jsonify({"error": "Login failed"}), 500

@app.route('/api/admin/logout', methods=['POST'])
def admin_logout():
    """Admin logout endpoint"""
    session.pop('admin_authenticated', None)
    return jsonify({"success": True, "message": "Logged out"})

@app.route('/api/admin/check', methods=['GET'])
def check_admin():
    """Check if user is authenticated"""
    return jsonify({
        "authenticated": session.get('admin_authenticated', False)
    })

# ============= EXISTING ROUTES =============

def clean_metadata_file():
    path = 'uploads/samples/metadata.json'
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            raw = f.read()
        cleaned = re.sub(r'[\x00-\x09\x0B\x0C\x0E-\x1F]', '', raw)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(cleaned)

clean_metadata_file()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_old_files():
    try:
        current_time = datetime.now().timestamp()
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if filename in ['samples', 'simulations', 'questions.json', 'metadata.json']:
                continue
            if os.path.isfile(filepath):
                file_age = current_time - os.path.getmtime(filepath)
                if file_age > 3600:
                    os.remove(filepath)
    except Exception as e:
        print(f"Cleanup error: {e}")

def get_audio_duration(file_path):
    audio = AudioSegment.from_file(file_path)
    return len(audio) / 1000.0

def convert_to_wav(input_path, output_path):
    audio = AudioSegment.from_file(input_path)
    audio = audio.set_frame_rate(16000).set_channels(1)
    audio.export(output_path, format="wav")
    return output_path

def transcribe_audio(file_path):
    try:
        with open(file_path, 'rb') as audio_file:
            transcription = groq_client.audio.transcriptions.create(
                file=("audio.wav", audio_file.read()),
                model="whisper-large-v3-turbo",
                response_format="json",
            )
        
        duration = get_audio_duration(file_path)
        transcript_text = transcription.text if hasattr(transcription, 'text') else str(transcription)
        
        return {
            "text": transcript_text,
            "words": [],
            "duration": duration
        }
    except Exception as e:
        print(f"Transcription error: {str(e)}")
        raise Exception(f"Transcription failed: {str(e)}")

def grade_speech(topic, transcript_data):
    transcript_text = transcript_data["text"]
    total_words = len(transcript_text.split())
    duration = transcript_data["duration"]
    words_per_minute = (total_words / duration * 60) if duration > 0 else 0
    
    prompt = f"""You are an expert English speaking examiner. Grade the following speech response based on this rubric:

**Rubric (Total: 2.0 points)**
1. Content (0.9/2.0 points)
   - Sufficiently address all requirements of the test question
   - Develop supporting ideas with relevant reasons and examples
   - Display a range of original and practical ideas

2. Accuracy (0.6/2.0 points)
   - Demonstrate a wide variety of vocabulary and grammatical structures
   - Make correct use of words, grammatical structures and linking devices
   - Demonstrate correct pronunciation with appropriate intonation

3. Delivery (0.5/2.0 points)
   - Maintain fluency throughout
   - Demonstrate effective use of presentation skills

**Topic/Question:** {topic}

**Speech Transcript:** {transcript_text}

**Speech Metrics:**
- Total words: {total_words}
- Duration: {duration:.1f} seconds
- Speaking pace: {words_per_minute:.0f} words/minute

**Instructions:**
1. Provide scores for each criterion (rounded to 2 decimal places)
2. Give detailed feedback for each criterion with specific examples from the transcript
3. Point out both strengths and areas for improvement
4. Generate a comprehensive sample 2.0/2.0 response to the same topic that would take approximately 5 minutes to speak (around 600-750 words). The sample should:
   - Start with "My question is... (if question number is provided), and the prompt is... Here is my response." and then answer the question fully
   - End with "This is the end of my speech. Thank you."
   - Be detailed and well-structured with clear introduction, body paragraphs, and conclusion
   - Include specific examples, explanations, and supporting details
   - Demonstrate sophisticated vocabulary and varied sentence structures
   - Show natural flow with appropriate transitions
   - Be comprehensive enough to fill a 5-minute speaking time
   - Be creative in the introduction to hook the listener's attention

Note: 
- Return feedback in bullet points when appropriate to maximize clarity (Strengths, Weaknesses, Suggestions)
- Grade at C2 level of the CEFR framework

**Return your response in this EXACT JSON format:**
{{
    "scores": {{
        "content": 0.00,
        "accuracy": 0.00,
        "delivery": 0.00,
        "total": 0.00
    }},
    "feedback": {{
        "content": "Detailed feedback with examples...",
        "accuracy": "Detailed feedback with examples...",
        "delivery": "Detailed feedback with examples..."
    }},
    "sample_response": "A complete 2.0/2.0 sample response to the topic..."
}}"""

    response = groq_client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3
    )
    
    result_text = response.choices[0].message.content
    
    if "```json" in result_text:
        result_text = result_text.split("```json")[1].split("```")[0].strip()
    elif "```" in result_text:
        result_text = result_text.split("```")[1].split("```")[0].strip()
    
    result_text = result_text.replace('"', '"').replace('"', '"')
    result_text = result_text.replace(''', "'").replace(''', "'")
    result_text = result_text.replace('—', '-').replace('–', '-')
    result_text = result_text.replace('\u2018', "'").replace('\u2019', "'")
    result_text = result_text.replace('\u201c', '"').replace('\u201d', '"')
    result_text = result_text.replace('\u2013', '-').replace('\u2014', '-')
    result_text = re.sub(r'[\u200b-\u200f\u202a-\u202e\u2060\uFEFF]', '', result_text)
    result_text = result_text.replace('\u202f', ' ')
    result_text = result_text.replace('\ufeff', '')
    result_text = result_text.replace('\u00A0', ' ')
    result_text = re.sub(r'[^\x00-\x7F]+', '', result_text)
    result_text = re.sub(r'[\x00-\x1F\x7F]', '', result_text)
    
    return json.loads(result_text)

def generate_docx(topic, transcript, grading_result):
    doc = Document()
    
    title = doc.add_heading('necs. - Speech Feedback Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph(f"Topic: {topic}")
    doc.add_paragraph()
    
    doc.add_heading('Score Summary', 1)
    scores = grading_result['scores']
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    score_data = [
        ('Content', f"{scores['content']}/0.9"),
        ('Accuracy', f"{scores['accuracy']}/0.6"),
        ('Delivery', f"{scores['delivery']}/0.5"),
        ('', ''),
        ('TOTAL SCORE', f"{scores['total']}/2.0")
    ]
    
    for i, (criterion, score) in enumerate(score_data):
        table.rows[i].cells[0].text = criterion
        table.rows[i].cells[1].text = score
        if i == 4:
            for cell in table.rows[i].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('Detailed Feedback', 1)
    feedback = grading_result['feedback']
    
    doc.add_heading('1. Content', 2)
    doc.add_paragraph(feedback['content'])
    
    doc.add_heading('2. Accuracy', 2)
    doc.add_paragraph(feedback['accuracy'])
    
    doc.add_heading('3. Delivery', 2)
    doc.add_paragraph(feedback['delivery'])
    
    doc.add_page_break()
    
    doc.add_heading('Your Speech Transcript', 1)
    doc.add_paragraph(transcript)
    
    doc.add_page_break()
    
    doc.add_heading('Sample 2.0/2.0 Response', 1)
    doc.add_paragraph(grading_result['sample_response'])
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

@app.route('/')
def serve():
    return send_from_directory(app.static_folder, 'index.html')

@app.errorhandler(404)
def not_found(e):
    if request.path.startswith('/api/'):
        return jsonify({"error": "API endpoint not found"}), 404
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/api', methods=['GET'])
def api_home():
    return jsonify({
        "message": "necs. API is running!",
        "version": "2.0",
        "security": "enabled"
    })

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy"})

@app.route('/api/analyze', methods=['POST'])
@rate_limit(max_requests=10, window_seconds=3600)
def analyze_speech():
    try:
        cleanup_old_files()
        
        if 'audio' not in request.files:
            return jsonify({"error": "No audio file provided"}), 400
        
        if 'topic' not in request.form:
            return jsonify({"error": "No topic provided"}), 400
        
        audio_file = request.files['audio']
        topic = request.form['topic']
        
        if audio_file.filename == '':
            return jsonify({"error": "No file selected"}), 400
        
        if not allowed_file(audio_file.filename):
            return jsonify({"error": "Invalid file format"}), 400
        
        filename = secure_filename(audio_file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        audio_file.save(filepath)
        
        duration = get_audio_duration(filepath)
        if duration > 320:
            os.remove(filepath)
            return jsonify({"error": "Audio file exceeds 5 minute limit"}), 400
        
        wav_filepath = filepath.rsplit('.', 1)[0] + '_compressed.wav'
        convert_to_wav(filepath, wav_filepath)
        if filepath != wav_filepath:
            os.remove(filepath)
        filepath = wav_filepath
        
        file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        
        if file_size_mb > 20:
            os.remove(filepath)
            return jsonify({"error": "Audio file too large"}), 400
        
        transcript_data = transcribe_audio(filepath)
        grading_result = grade_speech(topic, transcript_data)
        doc_stream = generate_docx(topic, transcript_data["text"], grading_result)
        
        os.remove(filepath)
        
        doc_bytes = doc_stream.getvalue()
        doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')
        
        return jsonify({
            "success": True,
            "transcript": transcript_data["text"],
            "duration": transcript_data["duration"],
            "scores": grading_result["scores"],
            "feedback": grading_result["feedback"],
            "sample_response": grading_result["sample_response"],
            "document_base64": doc_base64,
            "document_filename": f"necs_feedback_{timestamp}.docx"
        })
    
    except Exception as e:
        print(f"Error: {str(e)}")
        if 'filepath' in locals() and os.path.exists(filepath):
            os.remove(filepath)
        return jsonify({"error": str(e)}), 500

# ============= SECURED ADMIN ROUTES =============

@app.route('/api/samples', methods=['GET'])
def get_samples():
    try:
        samples = Sample.query.order_by(Sample.created_at.desc()).all()
        return jsonify({"samples": [s.to_dict() for s in samples]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/samples/upload', methods=['POST'])
@require_admin()
@rate_limit(max_requests=20, window_seconds=3600)
def upload_sample():
    try:
        if 'audio' not in request.files:
            return jsonify({"error": "No audio file"}), 400
        
        audio_file = request.files['audio']
        topic = request.form.get('topic')
        question = request.form.get('question', '')
        speaker = request.form.get('speaker')
        score = float(request.form.get('score', 2.0))
        transcript = request.form.get('transcript', '')
        feedback = request.form.get('feedback', '')
        
        if not all([topic, speaker, transcript, feedback]):
            return jsonify({"error": "Missing required fields"}), 400
        
        filename = secure_filename(audio_file.filename)
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        audio_file.save(temp_path)
        
        try:
            duration = int(get_audio_duration(temp_path))
        except:
            duration = 0
        
        upload_result = cloudinary.uploader.upload(
            temp_path,
            resource_type="video",
            folder="necs_samples",
            public_id=f"sample_{datetime.now().strftime('%Y%m%d_%H%M%S')}",
            overwrite=True
        )
        
        os.remove(temp_path)
        
        new_sample = Sample(
            filename=filename,
            topic=topic,
            question=question,
            speaker=speaker,
            score=score,
            duration=duration,
            transcript=transcript,
            feedback=feedback,
            audio_url=upload_result['secure_url']
        )
        
        db.session.add(new_sample)
        db.session.commit()
        
        return jsonify({"success": True, "id": new_sample.id})
        
    except Exception as e:
        db.session.rollback()
        if 'temp_path' in locals() and os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({"error": str(e)}), 500

@app.route('/api/samples/<int:sample_id>', methods=['PUT'])
@require_admin()
def update_sample(sample_id):
    try:
        sample = Sample.query.get(sample_id)
        if not sample:
            return jsonify({"error": "Not found"}), 404

        if 'topic' in request.form: sample.topic = request.form['topic']
        if 'question' in request.form: sample.question = request.form['question']
        if 'speaker' in request.form: sample.speaker = request.form['speaker']
        if 'score' in request.form: sample.score = float(request.form['score'])
        if 'transcript' in request.form: sample.transcript = request.form['transcript']
        if 'feedback' in request.form: sample.feedback = request.form['feedback']

        db.session.commit()
        return jsonify({"success": True})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/samples/<int:sample_id>', methods=['DELETE'])
@require_admin()
def delete_sample(sample_id):
    try:
        sample = Sample.query.get(sample_id)
        if not sample:
            return jsonify({"error": "Not found"}), 404

        db.session.delete(sample)
        db.session.commit()
        return jsonify({"success": True})
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/questions', methods=['GET'])
def get_questions():
    try:
        questions = Question.query.all()
        return jsonify({"questions": [q.to_dict() for q in questions]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/questions', methods=['POST'])
@require_admin()
def add_question():
    try:
        data = request.get_json()
        new_question = Question(
            topic=data['topic'],
            question=data['question'],
            category=data.get('category', 'General')
        )
        db.session.add(new_question)
        db.session.commit()
        return jsonify({"success": True, "id": new_question.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/questions/<int:question_id>', methods=['PUT'])
@require_admin()
def update_question(question_id):
    try:
        data = request.get_json()
        question = Question.query.get(question_id)
        if not question:
            return jsonify({"error": "Not found"}), 404
        
        if 'topic' in data: question.topic = data['topic']
        if 'question' in data: question.question = data['question']
        if 'category' in data: question.category = data['category']
        
        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/questions/<int:question_id>', methods=['DELETE'])
@require_admin()
def delete_question(question_id):
    try:
        question = Question.query.get(question_id)
        if not question:
            return jsonify({"error": "Not found"}), 404
        db.session.delete(question)
        db.session.commit()
        return jsonify({"success": True})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500

@app.route('/api/questions/random', methods=['GET'])
def get_random_question():
    try:
        questions = Question.query.all()
        if not questions:
            return jsonify({"error": "No questions"}), 404
        return jsonify({"question": random.choice(questions).to_dict()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)